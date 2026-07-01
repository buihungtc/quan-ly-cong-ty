from flask import Flask, jsonify, request, render_template, send_file
import json
import os
import sys
import webbrowser
import threading
import time
import tempfile
import database
import docx
import requests
from datetime import datetime
    
# Lấy thư mục gốc tương thích với PyInstaller & Electron
if os.environ.get('APP_BASE_DIR'):
    BASE_DIR = os.environ.get('APP_BASE_DIR')
elif getattr(sys, 'frozen', False):
    BASE_DIR = getattr(sys, '_MEIPASS', os.path.dirname(sys.executable))
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))

app = Flask(__name__,
            template_folder=os.path.join(BASE_DIR, 'templates'),
            static_folder=os.path.join(BASE_DIR, 'static'))

# Cấu hình mặc định
CONFIG_PATH = os.path.join(BASE_DIR, 'config.json')
DEFAULT_CONFIG = {
    "PASSCODE": "123456",
    "NGROK_AUTHTOKEN": "",
    "GEMINI_API_KEY": "",
    "TELEGRAM_ENABLED": False,
    "TELEGRAM_TIME": "09:00",
    "TELEGRAM_BOT_TOKEN": "",
    "TELEGRAM_CHAT_ID": ""
}

def load_config():
    """Tải cấu hình từ config.json"""
    if not os.path.exists(CONFIG_PATH):
        with open(CONFIG_PATH, 'w', encoding='utf-8') as f:
            json.dump(DEFAULT_CONFIG, f, indent=2, ensure_ascii=False)
        return DEFAULT_CONFIG
    try:
        with open(CONFIG_PATH, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception:
        return DEFAULT_CONFIG

# Load dotenv if available
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

config = load_config()

# Ưu tiên đọc từ biến môi trường (Docker/Server/.env), nếu không có thì đọc từ config.json (Local/Electron)
PASSCODE = os.environ.get("PASSCODE") or config.get("PASSCODE", "123456")
NGROK_AUTHTOKEN = os.environ.get("NGROK_AUTHTOKEN") or config.get("NGROK_AUTHTOKEN", "")
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY") or config.get("GEMINI_API_KEY", "")

# Xử lý kiểu boolean cho biến Telegram Enabled từ env
env_telegram_enabled = os.environ.get("TELEGRAM_ENABLED")
if env_telegram_enabled is not None:
    TELEGRAM_ENABLED = env_telegram_enabled.lower() in ("true", "1", "yes")
else:
    TELEGRAM_ENABLED = config.get("TELEGRAM_ENABLED", False)

TELEGRAM_TIME = os.environ.get("TELEGRAM_TIME") or config.get("TELEGRAM_TIME", "09:00")
TELEGRAM_BOT_TOKEN = os.environ.get("TELEGRAM_BOT_TOKEN") or config.get("TELEGRAM_BOT_TOKEN", "")
TELEGRAM_CHAT_ID = os.environ.get("TELEGRAM_CHAT_ID") or config.get("TELEGRAM_CHAT_ID", "")


# Khởi tạo Database
database.init_db()



# Biến lưu ngrok public URL và lỗi (được cập nhật bởi thread ngrok)
NGROK_PUBLIC_URL = None
NGROK_ERROR = None

# Helper lấy thông tin người dùng hiện tại từ token phiên làm việc
def get_current_user():
    token = request.headers.get('X-Auth-Token') or request.headers.get('X-Passcode')
    if not token:
        return None
    return database.verify_session(token)

# Helper kiểm tra quyền Admin cho các API ghi/chỉnh sửa dữ liệu (tương thích ngược tên hàm)
def verify_passcode():
    # Kiểm tra token phiên trước
    user = get_current_user()
    if user:
        return user.get('role') == 'admin'
    
    # Fallback kiểm tra passcode dạng thô
    client_passcode = request.headers.get('X-Passcode')
    if client_passcode == PASSCODE:
        return True
    
    # Kiểm tra nếu passcode thô chính là mật khẩu admin trong DB
    if client_passcode:
        user_db = database.verify_user('admin', client_passcode)
        if user_db:
            return True
            
    return False

# ----------------- ROUTES GIAO DIỆN -----------------
@app.route('/')
def index():
    return render_template('index.html')

# ----------------- ROUTES API -----------------
@app.route('/api/contracts', methods=['GET'])
def api_get_contracts():
    if not get_current_user():
        return jsonify({"error": "Vui lòng đăng nhập để xem dữ liệu."}), 401
    status = request.args.get('status') # 'processing' hoặc 'completed'
    try:
        contracts = database.get_contracts(status)
        return jsonify(contracts)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/parse-contract', methods=['POST'])
def api_parse_contract():
    if not verify_passcode():
        return jsonify({"error": "Chưa mở khóa hoặc Mã khóa không chính xác."}), 401
    
    if not GEMINI_API_KEY:
        return jsonify({"error": "Chưa cấu hình Claude API Key. Vui lòng cập nhật trong phần cài đặt."}), 400
        
    if 'file' not in request.files:
        return jsonify({"error": "Không tìm thấy file tải lên."}), 400
        
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "Chưa chọn file."}), 400
        
    if not file.filename.lower().endswith('.docx'):
        return jsonify({"error": "Vui lòng tải lên file Word (.docx)."}), 400
        
    try:
        # Save temp file
        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, file.filename)
        file.save(temp_path)
        
        # Read docx text
        doc = docx.Document(temp_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        text_content = '\\n'.join(full_text)
        
        # Clean up temp file
        if os.path.exists(temp_path):
            os.remove(temp_path)
            
        # Send to AI API (Vilao / Claude)
        prompt = """Bạn là trợ lý ảo đọc hợp đồng. Hãy đọc nội dung hợp đồng dưới đây và trả về DỮ LIỆU JSON CHUẨN XÁC, không kèm theo bất kỳ văn bản giải thích nào khác. Các trường cần trích xuất trong đối tượng JSON chính:
- contract_code: Số hợp đồng (hoặc mã hợp đồng).
- partner_name: Tên đối tác (bên A hoặc bên thuê/mua hoặc đối tác ký kết).
- start_date: Ngày bắt đầu thực hiện (định dạng YYYY-MM-DD, nếu không rõ có thể để trống).
- end_date: Ngày kết thúc hoặc thời hạn (định dạng YYYY-MM-DD, nếu không rõ có thể để trống).
- value: Tổng giá trị hợp đồng (số, không có dấu phẩy hay ký tự, ví dụ: 150000000).
- installments: Mảng các đợt thanh toán, mỗi phần tử gồm:
  - amount: số tiền đợt thanh toán (số)
  - deadline_date: thời hạn thanh toán (định dạng YYYY-MM-DD, nếu có)
- tasks: Mảng các công việc cần thực hiện trích xuất từ nội dung hợp đồng. Mỗi công việc gồm:
  - task_name: Tên công việc cụ thể (Ví dụ: "Hoàn thiện thủ tục giấy tờ ký hợp đồng", "Giục thanh toán đặt cọc 10%", "Làm biên bản nghiệm thu giai đoạn 1 và gửi đối tác", "Làm hóa đơn đợt 2 gửi đối tác", "Giục thanh toán tiền đợt cuối"). Tên công việc phải ghi rõ chi tiết ngắn gọn.
  - status_type: Số nguyên từ 1 đến 6 đại diện cho phân loại công việc theo các trạng thái thực tế sau:
    1: Chờ ký kết (Ví dụ: công việc liên quan đến đàm phán ký kết)
    2: Hoàn thiện thủ tục giấy tờ ký hợp đồng (Ví dụ: chuẩn bị hồ sơ, chuẩn bị phụ lục, giấy tờ liên quan để ký kết)
    3: Giục thanh toán hoặc đặt cọc nếu có (Ví dụ: giục đặt cọc, thanh toán đợt đầu/tạm ứng)
    4: Làm biên bản nghiệm thu và gửi biên bản nghiệm thu cho đối tác (Ví dụ: nghiệm thu bàn giao các giai đoạn)
    5: Làm hóa đơn gửi đối tác (Ví dụ: xuất hóa đơn tài chính VAT sau nghiệm thu)
    6: Giục thanh toán tiền (Ví dụ: nhắc nợ, đòi thanh toán đợt 2, 3, đợt cuối sau nghiệm thu hoặc hóa đơn)
  - target_date: Ngày hạn cần hoàn thành công việc (định dạng YYYY-MM-DD). Hãy tính toán logic dựa trên ngày bắt đầu hợp đồng (start_date) hoặc các mốc thời gian ghi trong hợp đồng (Ví dụ: thanh toán sau 7 ngày từ ngày ký, bàn giao thiết bị sau 15 ngày từ ngày ký...). Nếu không có mốc cụ thể, hãy ước tính hợp lý hoặc lấy mặc định là start_date.

Nội dung hợp đồng:
""" + text_content
        
        resp = requests.post(
            "https://api.vilao.ai/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {GEMINI_API_KEY}",
                "Content-Type": "application/json",
            },
            json={
                "model": "claude-opus-4-8",
                "messages": [
                    {"role": "user", "content": prompt}
                ],
                "stream": False,
            },
            timeout=120,
        )
        resp.raise_for_status()
        data = resp.json()
        resp_text = data["choices"][0]["message"]["content"]
        
        # Parse JSON
        if "```json" in resp_text:
            resp_text = resp_text.split("```json")[1].split("```")[0]
        elif "```" in resp_text:
            resp_text = resp_text.split("```")[1].split("```")[0]
        else:
            start = resp_text.find('{')
            end = resp_text.rfind('}')
            if start != -1 and end != -1:
                resp_text = resp_text[start:end+1]
                
        parsed_data = json.loads(resp_text.strip())
        return jsonify({"success": True, "data": parsed_data})
        
    except requests.exceptions.RequestException as e:
        return jsonify({"error": f"Lỗi gọi API AI: {str(e)}"}), 500
    except Exception as e:
        return jsonify({"error": f"Lỗi xử lý file hoặc AI: {str(e)}"}), 500

@app.route('/api/parse-contract-text', methods=['POST'])
def api_parse_contract_text():
    if not verify_passcode():
        return jsonify({"error": "Chưa mở khóa hoặc Mã khóa không chính xác."}), 401
        
    if not GEMINI_API_KEY:
        return jsonify({"error": "Chưa cấu hình Claude API Key. Vui lòng cập nhật trong phần cài đặt."}), 400
        
    data = request.json or {}
    text_content = data.get('contract_text', '').strip()
    if not text_content:
        return jsonify({"error": "Nội dung hợp đồng trống."}), 400
        
    try:
        prompt = """Bạn là trợ lý ảo đọc hợp đồng. Hãy đọc nội dung hợp đồng dưới đây và trả về DỮ LIỆU JSON CHUẨN XÁC, không kèm theo bất kỳ văn bản giải thích nào khác. Các trường cần trích xuất trong đối tượng JSON chính:
- contract_code: Số hợp đồng (hoặc mã hợp đồng).
- partner_name: Tên đối tác (bên A hoặc bên thuê/mua hoặc đối tác ký kết).
- start_date: Ngày bắt đầu thực hiện (định dạng YYYY-MM-DD, nếu không rõ có thể để trống).
- end_date: Ngày kết thúc hoặc thời hạn (định dạng YYYY-MM-DD, nếu không rõ có thể để trống).
- value: Tổng giá trị hợp đồng (số, không có dấu phẩy hay ký tự, ví dụ: 150000000).
- installments: Mảng các đợt thanh toán, mỗi phần tử gồm:
  - amount: số tiền đợt thanh toán (số)
  - deadline_date: thời hạn thanh toán (định dạng YYYY-MM-DD, nếu có)
- tasks: Mảng các công việc cần thực hiện trích xuất từ nội dung hợp đồng. Mỗi công việc gồm:
  - task_name: Tên công việc cụ thể (Ví dụ: "Hoàn thiện thủ tục giấy tờ ký hợp đồng", "Giục thanh toán đặt cọc 10%", "Làm biên bản nghiệm thu giai đoạn 1 và gửi đối tác", "Làm hóa đơn đợt 2 gửi đối tác", "Giục thanh toán tiền đợt cuối"). Tên công việc phải ghi rõ chi tiết ngắn gọn.
  - status_type: Số nguyên từ 1 đến 6 đại diện cho phân loại công việc theo các trạng thái thực tế sau:
    1: Chờ ký kết (Ví dụ: công việc liên quan đến đàm phán ký kết)
    2: Hoàn thiện thủ tục giấy tờ ký hợp đồng (Ví dụ: chuẩn bị hồ sơ, chuẩn bị phụ lục, giấy tờ liên quan để ký kết)
    3: Giục thanh toán hoặc đặt cọc nếu có (Ví dụ: giục đặt cọc, thanh toán đợt đầu/tạm ứng)
    4: Làm biên bản nghiệm thu và gửi biên bản nghiệm thu cho đối tác (Ví dụ: nghiệm thu bàn giao các giai đoạn)
    5: Làm hóa đơn gửi đối tác (Ví dụ: xuất hóa đơn tài chính VAT sau nghiệm thu)
    6: Giục thanh toán tiền (Ví dụ: nhắc nợ, đòi thanh toán đợt 2, 3, đợt cuối sau nghiệm thu hoặc hóa đơn)
  - target_date: Ngày hạn cần hoàn thành công việc (định dạng YYYY-MM-DD). Hãy tính toán logic dựa trên ngày bắt đầu hợp đồng (start_date) hoặc các mốc thời gian ghi trong hợp đồng (Ví dụ: thanh toán sau 7 ngày từ ngày ký, bàn giao thiết bị sau 15 ngày từ ngày ký...). Nếu không có mốc cụ thể, hãy ước tính hợp lý hoặc lấy mặc định là start_date.

Nội dung hợp đồng:
""" + text_content

        resp = requests.post(
            "https://api.vilao.ai/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {GEMINI_API_KEY}",
                "Content-Type": "application/json",
            },
            json={
                "model": "claude-opus-4-8",
                "messages": [{"role": "user", "content": prompt}],
                "stream": False,
            },
            timeout=120,
        )
        resp.raise_for_status()
        resp_data = resp.json()
        resp_text = resp_data["choices"][0]["message"]["content"]
        
        # Parse JSON
        if "```json" in resp_text:
            resp_text = resp_text.split("```json")[1].split("```")[0]
        elif "```" in resp_text:
            resp_text = resp_text.split("```")[1].split("```")[0]
        else:
            start = resp_text.find('{')
            end = resp_text.rfind('}')
            if start != -1 and end != -1:
                resp_text = resp_text[start:end+1]
                
        parsed_data = json.loads(resp_text.strip())
        return jsonify({"success": True, "data": parsed_data})
    except Exception as e:
        return jsonify({"error": f"Lỗi phân tích văn bản: {str(e)}"}), 500

@app.route('/api/contracts', methods=['POST'])
def api_add_contract():
    if not verify_passcode():
        return jsonify({"error": "Chưa mở khóa hoặc Mã khóa không chính xác. Bạn không có quyền thêm mới."}), 401
    
    data = request.json
    if not data:
        return jsonify({"error": "Dữ liệu không hợp lệ."}), 400
        
    contract_code = data.get('contract_code', '').strip()
    partner_name = data.get('partner_name', '').strip()
    value = data.get('value', 0)
    start_date = data.get('start_date', '').strip()
    end_date = data.get('end_date', '').strip()
    progress_notes = data.get('progress_notes', '').strip()
    installments_data = data.get('installments', [])
    tasks_data = data.get('tasks', [])
    
    bank_account_id = data.get('bank_account_id')
    if bank_account_id == '':
        bank_account_id = None
    elif bank_account_id is not None:
        try:
            bank_account_id = int(bank_account_id)
        except ValueError:
            bank_account_id = None
        
    try:
        value = float(value) if value else 0
    except ValueError:
        value = 0
        
    contract_id, error_msg = database.add_contract(
        contract_code, partner_name, value, start_date, end_date, progress_notes, installments_data, tasks_data, bank_account_id
    )
    
    if error_msg:
        return jsonify({"error": error_msg}), 400
        
    return jsonify({"message": "Thêm hợp đồng mới thành công!", "id": contract_id}), 201

@app.route('/api/contracts/<int:contract_id>', methods=['GET'])
def api_get_contract(contract_id):
    if not get_current_user():
        return jsonify({"error": "Vui lòng đăng nhập để xem dữ liệu."}), 401
    try:
        contract = database.get_contract_by_id(contract_id)
        if contract:
            return jsonify(contract)
        else:
            return jsonify({"error": "Không tìm thấy hợp đồng."}), 404
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/contracts/<int:contract_id>', methods=['PUT'])
def api_edit_contract(contract_id):
    if not verify_passcode():
        return jsonify({"error": "Chưa mở khóa hoặc Mã khóa không chính xác. Bạn không có quyền sửa."}), 401
    
    data = request.json
    if not data:
        return jsonify({"error": "Dữ liệu không hợp lệ."}), 400
        
    contract_code = data.get('contract_code', '').strip()
    partner_name = data.get('partner_name', '').strip()
    value = data.get('value', 0)
    start_date = data.get('start_date', '').strip()
    end_date = data.get('end_date', '').strip()
    progress_notes = data.get('progress_notes', '').strip()
    installments_data = data.get('installments', [])
    tasks_data = data.get('tasks', [])
    
    bank_account_id = data.get('bank_account_id')
    if bank_account_id == '':
        bank_account_id = None
    elif bank_account_id is not None:
        try:
            bank_account_id = int(bank_account_id)
        except ValueError:
            bank_account_id = None
        
    try:
        value = float(value) if value else 0
    except ValueError:
        value = 0
        
    success, error_msg = database.update_contract(
        contract_id, contract_code, partner_name, value, start_date, end_date, progress_notes, installments_data, tasks_data, bank_account_id
    )
    
    if not success:
        return jsonify({"error": error_msg or "Lỗi cập nhật hợp đồng."}), 400
        
    return jsonify({"message": "Cập nhật hợp đồng thành công!"}), 200

@app.route('/api/contracts/<int:contract_id>/notes', methods=['PUT'])
def api_update_notes(contract_id):
    if not verify_passcode():
        return jsonify({"error": "Chưa mở khóa hoặc Mã khóa không chính xác. Bạn không có quyền chỉnh sửa."}), 401
        
    data = request.json
    if not data or 'progress_notes' not in data:
        return jsonify({"error": "Thiếu dữ liệu ghi chú tiến độ."}), 400
        
    progress_notes = data.get('progress_notes', '').strip()
    
    success = database.update_progress_notes(contract_id, progress_notes)
    if success:
        return jsonify({"message": "Cập nhật ghi chú tiến độ thành công!"})
    else:
        return jsonify({"error": "Không tìm thấy hợp đồng hoặc không có thay đổi."}), 404

@app.route('/api/contracts/<int:contract_id>/complete', methods=['PUT'])
def api_mark_complete(contract_id):
    if not verify_passcode():
        return jsonify({"error": "Chưa mở khóa hoặc Mã khóa không chính xác. Bạn không có quyền cập nhật trạng thái."}), 401
        
    success = database.mark_contract_completed(contract_id)
    if success:
        return jsonify({"message": "Đã đánh dấu hoàn thành hợp đồng!"})
    else:
        return jsonify({"error": "Không tìm thấy hợp đồng."}), 404

@app.route('/api/contracts/<int:contract_id>/revert', methods=['PUT'])
def api_revert_complete(contract_id):
    if not verify_passcode():
        return jsonify({"error": "Chưa mở khóa hoặc Mã khóa không chính xác. Bạn không có quyền hoàn tác trạng thái."}), 401
        
    success = database.revert_contract_status(contract_id)
    if success:
        return jsonify({"message": "Đã hoàn tác hợp đồng về trạng thái Đang xử lý!"})
    else:
        return jsonify({"error": "Không tìm thấy hợp đồng."}), 404

@app.route('/api/contracts/<int:contract_id>', methods=['DELETE'])
def api_delete_contract(contract_id):
    if not verify_passcode():
        return jsonify({"error": "Chưa mở khóa hoặc Mã khóa không chính xác. Bạn không có quyền xóa hợp đồng."}), 401
        
    success = database.delete_contract(contract_id)
    if success:
        return jsonify({"message": "Đã xóa hợp đồng thành công!"})
    else:
        return jsonify({"error": "Không tìm thấy hợp đồng."}), 404

@app.route('/api/installments/<int:installment_id>/pay', methods=['PUT'])
def api_pay_installment(installment_id):
    if not verify_passcode():
        return jsonify({"error": "Chưa mở khóa hoặc Mã khóa không chính xác. Bạn không có quyền chỉnh sửa."}), 401
        
    data = request.json or {}
    is_paid = data.get('is_paid', False)
    paid_date = data.get('paid_date', '')
    
    success = database.update_installment_payment(installment_id, is_paid, paid_date)
    if success:
        return jsonify({"message": "Cập nhật thanh toán thành công!"})
    else:
        return jsonify({"error": "Không tìm thấy đợt thanh toán."}), 404

@app.route('/api/tasks/<int:task_id>/complete', methods=['PUT'])
def api_complete_task(task_id):
    if not verify_passcode():
        return jsonify({"error": "Chưa mở khóa hoặc Mã khóa không chính xác."}), 401
    
    data = request.json or {}
    is_completed = data.get('is_completed', False)
    completed_date = datetime.now().strftime('%Y-%m-%d') if is_completed else None
    
    success = database.update_task_status(task_id, is_completed, completed_date)
    if success:
        return jsonify({"message": "Cập nhật trạng thái công việc thành công!"})
    else:
        return jsonify({"error": "Không tìm thấy công việc."}), 404

@app.route('/api/stats', methods=['GET'])
def api_get_stats():
    if not get_current_user():
        return jsonify({"error": "Vui lòng đăng nhập để xem dữ liệu."}), 401
    try:
        bank_account_id = request.args.get('bank_account_id')
        stats = database.get_statistics(bank_account_id)
        return jsonify(stats)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/login', methods=['POST'])
def api_login():
    data = request.json or {}
    username = data.get('username', '').strip()
    password = data.get('password', '').strip()
    
    if not username or not password:
        return jsonify({"success": False, "error": "Thiếu tên đăng nhập hoặc mật khẩu."}), 400
        
    user = database.verify_user(username, password)
    if user:
        token = database.generate_token(user['username'], user['role'])
        return jsonify({
            "success": True,
            "username": user['username'],
            "role": user['role'],
            "token": token
        })
    else:
        return jsonify({"success": False, "error": "Tên đăng nhập hoặc mật khẩu không chính xác."}), 401

@app.route('/api/logout', methods=['POST'])
def api_logout():
    token = request.headers.get('X-Auth-Token')
    if token:
        database.delete_session(token)
    return jsonify({"success": True, "message": "Đã đăng xuất thành công."})

@app.route('/api/verify-token', methods=['GET'])
def api_verify_token():
    token = request.headers.get('X-Auth-Token')
    if not token:
        return jsonify({"success": False, "error": "Thiếu token xác thực."}), 401
    user = database.verify_session(token)
    if user:
        return jsonify({
            "success": True,
            "username": user['username'],
            "role": user['role']
        })
    else:
        return jsonify({"success": False, "error": "Phiên đăng nhập không hợp lệ hoặc đã hết hạn."}), 401

@app.route('/api/admin/change-password', methods=['POST'])
def api_change_password():
    user = get_current_user()
    if not user or user.get('role') != 'admin':
        return jsonify({"error": "Bạn không có quyền thực hiện hành động này."}), 403
        
    data = request.json or {}
    new_password = data.get('new_password', '').strip()
    if not new_password:
        return jsonify({"error": "Mật khẩu mới không được để trống."}), 400
        
    success = database.update_user_password('admin', new_password)
    if success:
        # Đồng bộ passcode cũ để tương thích với Electron
        global PASSCODE
        PASSCODE = new_password
        try:
            cfg = load_config()
            cfg['PASSCODE'] = new_password
            with open(CONFIG_PATH, 'w', encoding='utf-8') as f:
                json.dump(cfg, f, indent=2, ensure_ascii=False)
        except Exception:
            pass
        return jsonify({"success": True, "message": "Đổi mật khẩu Admin thành công!"})
    else:
        return jsonify({"error": "Không thể đổi mật khẩu."}), 500

@app.route('/api/admin/guests', methods=['GET'])
def api_get_guests():
    user = get_current_user()
    if not user or user.get('role') != 'admin':
        return jsonify({"error": "Bạn không có quyền thực hiện hành động này."}), 403
    guests = database.get_guest_users()
    return jsonify(guests)

@app.route('/api/admin/create-guest', methods=['POST'])
def api_create_guest():
    user = get_current_user()
    if not user or user.get('role') != 'admin':
        return jsonify({"error": "Bạn không có quyền thực hiện hành động này."}), 403
        
    data = request.json or {}
    username = data.get('username', '').strip()
    password = data.get('password', '').strip()
    
    if not username or not password:
        return jsonify({"error": "Thiếu tên đăng nhập hoặc mật khẩu tài khoản khách."}), 400
        
    if username == 'admin':
        return jsonify({"error": "Không thể tạo tài khoản khách trùng tên với admin."}), 400
        
    success, err = database.create_guest_user(username, password)
    if success:
        return jsonify({"success": True, "message": f"Tạo tài khoản khách '{username}' thành công!"}), 201
    else:
        return jsonify({"error": err or "Không thể tạo tài khoản khách."}), 400

@app.route('/api/admin/guests/<username>', methods=['DELETE'])
def api_delete_guest(username):
    user = get_current_user()
    if not user or user.get('role') != 'admin':
        return jsonify({"error": "Bạn không có quyền thực hiện hành động này."}), 403
        
    if username == 'guest':
        return jsonify({"error": "Không thể xóa tài khoản khách mặc định."}), 400
        
    success = database.delete_guest_user(username)
    if success:
        return jsonify({"success": True, "message": f"Đã xóa tài khoản khách '{username}'."})
    else:
        return jsonify({"error": "Không tìm thấy tài khoản khách."}), 404

# ----------------- ROUTES API TÀI KHOẢN NGÂN HÀNG THỤ HƯỞNG -----------------
@app.route('/api/bank-accounts', methods=['GET'])
def api_get_bank_accounts():
    if not get_current_user():
        return jsonify({"error": "Vui lòng đăng nhập để xem dữ liệu."}), 401
    try:
        accounts = database.get_bank_accounts()
        return jsonify(accounts)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/bank-accounts', methods=['POST'])
def api_add_bank_account():
    if not verify_passcode():
        return jsonify({"error": "Bạn không có quyền thực hiện hành động này."}), 403
    
    data = request.json or {}
    bank_name = data.get('bank_name', '').strip()
    account_number = data.get('account_number', '').strip()
    account_holder = data.get('account_holder', '').strip()
    description = data.get('description', '').strip()
    
    if not bank_name or not account_number or not account_holder:
        return jsonify({"error": "Vui lòng nhập đầy đủ Tên ngân hàng, Số tài khoản và Tên chủ tài khoản."}), 400
        
    account_id, error_msg = database.add_bank_account(bank_name, account_number, account_holder, description)
    if error_msg:
        return jsonify({"error": error_msg}), 400
        
    return jsonify({"message": "Thêm tài khoản ngân hàng thụ hưởng thành công!", "id": account_id}), 201

@app.route('/api/bank-accounts/<int:bank_account_id>', methods=['PUT'])
def api_update_bank_account(bank_account_id):
    if not verify_passcode():
        return jsonify({"error": "Bạn không có quyền thực hiện hành động này."}), 403
        
    data = request.json or {}
    bank_name = data.get('bank_name', '').strip()
    account_number = data.get('account_number', '').strip()
    account_holder = data.get('account_holder', '').strip()
    description = data.get('description', '').strip()
    
    if not bank_name or not account_number or not account_holder:
        return jsonify({"error": "Vui lòng nhập đầy đủ Tên ngân hàng, Số tài khoản và Tên chủ tài khoản."}), 400
        
    success, error_msg = database.update_bank_account(bank_account_id, bank_name, account_number, account_holder, description)
    if error_msg:
        return jsonify({"error": error_msg}), 400
        
    if success:
        return jsonify({"message": "Cập nhật tài khoản ngân hàng thụ hưởng thành công!"})
    else:
        return jsonify({"error": "Không tìm thấy tài khoản ngân hàng hoặc số tài khoản đã tồn tại."}), 404

@app.route('/api/bank-accounts/<int:bank_account_id>', methods=['DELETE'])
def api_delete_bank_account(bank_account_id):
    if not verify_passcode():
        return jsonify({"error": "Bạn không có quyền thực hiện hành động này."}), 403
        
    success = database.delete_bank_account(bank_account_id)
    if success:
        return jsonify({"message": "Đã xóa tài khoản ngân hàng thụ hưởng thành công!"})
    else:
        return jsonify({"error": "Không tìm thấy tài khoản ngân hàng."}), 404

@app.route('/api/verify-passcode', methods=['POST'])
def api_verify_passcode():
    data = request.json
    if not data or 'passcode' not in data:
        return jsonify({"success": False, "error": "Thiếu mã khóa."}), 400
    
    passcode = data.get('passcode')
    user = database.verify_user('admin', passcode)
    if passcode == PASSCODE or user is not None:
        return jsonify({"success": True})
    else:
        return jsonify({"success": False, "error": "Mã khóa không chính xác."})

@app.route('/api/server-info', methods=['GET'])
def api_server_info():
    """Trả về thông tin URL server để Electron hiển thị"""
    port = int(os.environ.get('FLASK_PORT', 5000))
    return jsonify({
        "local_url": f"http://127.0.0.1:{port}",
        "ngrok_url": NGROK_PUBLIC_URL,
        "ngrok_error": NGROK_ERROR
    })

@app.route('/api/ngrok/connect', methods=['POST'])
def api_ngrok_connect():
    """Kết nối lại ngrok, tùy chọn cập nhật authtoken mới"""
    global NGROK_PUBLIC_URL, NGROK_ERROR, NGROK_AUTHTOKEN
    data = request.json or {}
    new_token = str(data.get('authtoken', '')).strip()
    if new_token:
        NGROK_AUTHTOKEN = new_token
        try:
            current_cfg = load_config()
            current_cfg['NGROK_AUTHTOKEN'] = new_token
            with open(CONFIG_PATH, 'w', encoding='utf-8') as f:
                json.dump(current_cfg, f, indent=2, ensure_ascii=False)
        except Exception as e:
            return jsonify({"success": False, "error": f"Lỗi lưu config: {e}"}), 500
    NGROK_PUBLIC_URL = None
    NGROK_ERROR = None
    try:
        from pyngrok import ngrok as _ngrok
        _ngrok.kill()
    except Exception:
        pass
    port = int(os.environ.get('FLASK_PORT', 5000))
    threading.Thread(target=start_ngrok, args=(port,), daemon=True).start()
    return jsonify({"success": True, "message": "Đang kết nối lại ngrok..."})

@app.route('/api/config', methods=['GET', 'POST'])
def api_config():
    global PASSCODE, NGROK_AUTHTOKEN, NGROK_PUBLIC_URL, GEMINI_API_KEY
    global TELEGRAM_ENABLED, TELEGRAM_TIME, TELEGRAM_BOT_TOKEN, TELEGRAM_CHAT_ID
    if request.method == 'GET':
        return jsonify({
            "passcode": PASSCODE,
            "ngrok_token": NGROK_AUTHTOKEN,
            "gemini_api_key": GEMINI_API_KEY,
            "telegram_enabled": TELEGRAM_ENABLED,
            "telegram_time": TELEGRAM_TIME,
            "telegram_bot_token": TELEGRAM_BOT_TOKEN,
            "telegram_chat_id": TELEGRAM_CHAT_ID
        })
    elif request.method == 'POST':
        data = request.json or {}
        new_passcode = str(data.get('passcode', PASSCODE)).strip()
        new_ngrok_token = str(data.get('ngrok_token', data.get('ngrok_authtoken', NGROK_AUTHTOKEN))).strip()
        new_gemini_key = str(data.get('gemini_api_key', GEMINI_API_KEY)).strip()
        new_telegram_enabled = bool(data.get('telegram_enabled', TELEGRAM_ENABLED))
        new_telegram_time = str(data.get('telegram_time', TELEGRAM_TIME)).strip()
        new_telegram_bot_token = str(data.get('telegram_bot_token', TELEGRAM_BOT_TOKEN)).strip()
        new_telegram_chat_id = str(data.get('telegram_chat_id', TELEGRAM_CHAT_ID)).strip()

        # Update variables
        old_token = NGROK_AUTHTOKEN
        PASSCODE = new_passcode
        NGROK_AUTHTOKEN = new_ngrok_token
        GEMINI_API_KEY = new_gemini_key
        TELEGRAM_ENABLED = new_telegram_enabled
        TELEGRAM_TIME = new_telegram_time
        TELEGRAM_BOT_TOKEN = new_telegram_bot_token
        TELEGRAM_CHAT_ID = new_telegram_chat_id

        # Ghi cấu hình mới vào config.json
        try:
            with open(CONFIG_PATH, 'w', encoding='utf-8') as f:
                json.dump({
                    "PASSCODE": PASSCODE,
                    "NGROK_AUTHTOKEN": NGROK_AUTHTOKEN,
                    "GEMINI_API_KEY": GEMINI_API_KEY,
                    "TELEGRAM_ENABLED": TELEGRAM_ENABLED,
                    "TELEGRAM_TIME": TELEGRAM_TIME,
                    "TELEGRAM_BOT_TOKEN": TELEGRAM_BOT_TOKEN,
                    "TELEGRAM_CHAT_ID": TELEGRAM_CHAT_ID
                }, f, indent=2, ensure_ascii=False)
        except Exception as e:
            return jsonify({"success": False, "error": f"Lỗi ghi file config: {e}"}), 500

        # Nếu token thay đổi, khởi động lại ngrok tunnel ngay lập tức
        if new_ngrok_token != old_token:
            NGROK_PUBLIC_URL = None
            try:
                from pyngrok import ngrok
                # Hủy ngrok cũ
                ngrok.kill()
            except Exception:
                pass
            
            if new_ngrok_token:
                port = int(os.environ.get('FLASK_PORT', 5000))
                threading.Thread(target=start_ngrok, args=(port,), daemon=True).start()

        return jsonify({"success": True, "message": "Đã cập nhật và lưu cấu hình thành công!"})


def evaluate_contracts_with_ai(contracts):
    if not GEMINI_API_KEY:
        return None
        
    today_str = datetime.now().strftime("%d/%m/%Y")
    
    results = {}
    chunk_size = 10
    for i in range(0, len(contracts), chunk_size):
        chunk = contracts[i:i + chunk_size]
        
        prompt = f"""Bạn là trợ lý AI phân tích tiến độ hợp đồng. Hôm nay là ngày {today_str}.
Dưới đây là danh sách các hợp đồng đang thực hiện (định dạng JSON):
{json.dumps(chunk, ensure_ascii=False)}

Hãy phân tích từng hợp đồng và quyết định xem hợp đồng đó có CẦN NHẮC NHỞ HÔM NAY không.
Quy tắc:
1. NẾU trong progress_notes có dặn dò nhắc nhở (VD: "30/06 nhắc khách", "khất tới tháng sau"). Nếu chưa tới ngày hẹn -> KHÔNG cần nhắc nhở. NẾU đã tới hoặc qua ngày hẹn -> CẦN nhắc nhở.
2. NẾU hợp đồng đã quá hạn (end_date < hôm nay) HOẶC có đợt thanh toán chưa đóng tiền bị quá hạn -> Cần nhắc nhở.
3. NẾU mọi thứ bình thường, chưa tới hạn -> KHÔNG cần nhắc nhở.

Trả về KẾT QUẢ DƯỚI DẠNG JSON MẢNG (Array of objects), không giải thích thêm:
[
  {{"id": <contract_id>, "needs_attention": true/false, "suggestion": "Lý do ngắn gọn nếu cần nhắc nhở"}}
]
"""
        try:
            resp = requests.post(
                "https://api.vilao.ai/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {GEMINI_API_KEY}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": "claude-opus-4-8",
                    "messages": [{"role": "user", "content": prompt}],
                    "stream": False,
                },
                timeout=60,
            )
            if resp.status_code == 200:
                data = resp.json()
                resp_text = data["choices"][0]["message"]["content"]
                
                # Parse JSON
                if "```json" in resp_text:
                    resp_text = resp_text.split("```json")[1].split("```")[0]
                elif "```" in resp_text:
                    resp_text = resp_text.split("```")[1].split("```")[0]
                    
                parsed = json.loads(resp_text)
                for item in parsed:
                    results[item['id']] = item
            else:
                return None
        except Exception as e:
            print(f"[AI Telegram] Evaluate error: {e}")
            return None # Fallback nếu gọi AI lỗi
            
    return results

def send_telegram_report():
    try:
        contracts = database.get_contracts('processing')
        
        if not contracts:
            text = "📊 Báo cáo Hợp đồng hàng ngày\n\nKhông có hợp đồng nào đang thực hiện."
            url = f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendMessage"
            requests.post(url, json={"chat_id": TELEGRAM_CHAT_ID, "text": text}, timeout=10)
            return
            
        today_str = datetime.now().strftime('%Y-%m-%d')
        due_tasks_by_contract = []
        
        # Ánh xạ status_type sang tiếng Việt
        status_labels = {
            1: "Chờ ký kết",
            2: "Hoàn thiện giấy tờ",
            3: "Giục thanh toán/đặt cọc",
            4: "Làm & gửi biên bản nghiệm thu",
            5: "Làm hóa đơn gửi đối tác",
            6: "Giục thanh toán tiền"
        }
        
        for c in contracts:
            due_tasks = []
            for t in c.get('tasks', []):
                if not t.get('is_completed') and t.get('target_date'):
                    if t.get('target_date') <= today_str:
                        due_tasks.append(t)
            
            if due_tasks:
                due_tasks_by_contract.append((c, due_tasks))
                
        if not due_tasks_by_contract:
            text = "✅ Mọi thứ đều ổn, không có công việc hợp đồng nào cần xử lý khẩn cấp hôm nay."
        else:
            text = f"🚨 CÓ KHẨN CẤP: {len(due_tasks_by_contract)} HỢP ĐỒNG CÓ VIỆC CẦN XỬ LÝ HÔM NAY:\n\n"
            for c, tasks in due_tasks_by_contract:
                text += f"🔹 HĐ: {c['contract_code']} - {c['partner_name']}\n"
                for t in tasks:
                    status_name = status_labels.get(t['status_type'], "Khác")
                    try:
                        date_obj = datetime.strptime(t['target_date'], '%Y-%m-%d')
                        date_formatted = date_obj.strftime('%d/%m/%Y')
                    except Exception:
                        date_formatted = t['target_date']
                    text += f"   - Hạn: {date_formatted} | [{status_name}] {t['task_name']}\n"
                text += "\n"
        
        url = f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendMessage"
        requests.post(url, json={
            "chat_id": TELEGRAM_CHAT_ID,
            "text": text
        }, timeout=10)
    except Exception as e:
        print(f"[Telegram] Send report error: {e}")

def telegram_worker():
    last_sent_date = None
    while True:
        try:
            if TELEGRAM_ENABLED and TELEGRAM_BOT_TOKEN and TELEGRAM_CHAT_ID and TELEGRAM_TIME:
                now = datetime.now()
                if now.strftime('%H:%M') == TELEGRAM_TIME:
                    today = now.strftime('%Y-%m-%d')
                    if last_sent_date != today:
                        send_telegram_report()
                        last_sent_date = today
        except Exception as e:
            print(f"[Telegram] Worker error: {e}")
        time.sleep(30)

def start_ngrok(port):
    """Khởi động ngrok tunnel và lưu URL vào biến toàn cục"""
    global NGROK_PUBLIC_URL, NGROK_ERROR
    NGROK_ERROR = None
    if not NGROK_AUTHTOKEN:
        print("[ngrok] NGROK_AUTHTOKEN is not set or empty. Skipping ngrok setup.")
        return
    try:
        from pyngrok import ngrok
        ngrok.set_auth_token(NGROK_AUTHTOKEN)
        print("[ngrok] Authtoken applied.")
        public_url = ngrok.connect(port).public_url
        NGROK_PUBLIC_URL = public_url
        NGROK_ERROR = None

        print(f"\n========================================================")
        print(f"🚀 ngrok public URL: {public_url}")
        print(f"========================================================\n")
    except ImportError:
        NGROK_ERROR = "pyngrok chưa được cài đặt."
        print("[ngrok] pyngrok not installed. Running local only.")
    except Exception as e:
        err_str = str(e)
        if 'ERR_NGROK_108' in err_str or 'simultaneous' in err_str:
            NGROK_ERROR = "Tài khoản ngrok đang đạt giới hạn số phiên kết nối đồng thời."
        elif 'authentication failed' in err_str or 'ERR_NGROK_4' in err_str:
            NGROK_ERROR = "Xác thực thất bại: Authtoken không hợp lệ hoặc chưa đăng nhập ngrok."
        elif 'ERR_NGROK' in err_str:
            import re
            code = re.search(r'ERR_NGROK_\d+', err_str)
            NGROK_ERROR = f"Lỗi ngrok: {code.group() if code else 'Không xác định'}."
        else:
            NGROK_ERROR = f"Lỗi kết nối: {err_str[:100]}"
        print(f"[ngrok] Tunnel error: {e}")

@app.route('/api/ai-writer/templates', methods=['GET'])
def api_get_templates():
    if not get_current_user():
        return jsonify({"error": "Vui lòng đăng nhập để xem dữ liệu."}), 401
    templates = database.get_templates()
    return jsonify(templates)

@app.route('/api/ai-writer/templates', methods=['POST'])
def api_add_template():
    if not verify_passcode():
        return jsonify({"error": "Chưa mở khóa hoặc Mã khóa không chính xác."}), 401
        
    name = request.form.get('name', '').strip()
    content = request.form.get('content', '').strip()
    
    if not name:
        return jsonify({"error": "Tên mẫu hợp đồng không được bỏ trống."}), 400
        
    # Xử lý file docx đính kèm nếu có
    if 'file' in request.files:
        file = request.files['file']
        if file and file.filename != '':
            if not file.filename.lower().endswith('.docx'):
                return jsonify({"error": "Vui lòng đính kèm file Word (.docx)."}), 400
            try:
                temp_dir = tempfile.gettempdir()
                temp_path = os.path.join(temp_dir, file.filename)
                file.save(temp_path)
                
                doc = docx.Document(temp_path)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                content = '\n'.join(full_text)
                
                if os.path.exists(temp_path):
                    os.remove(temp_path)
            except Exception as e:
                return jsonify({"error": f"Lỗi đọc file docx: {str(e)}"}), 500
                
    if not content:
        return jsonify({"error": "Nội dung mẫu hợp đồng không được bỏ trống (hoặc phải đính kèm file docx)."}), 400
        
    template_id, err = database.add_template(name, content)
    if err:
        return jsonify({"error": err}), 400
        
    return jsonify({"success": True, "template_id": template_id, "message": "Thêm mẫu hợp đồng thành công."}), 201

@app.route('/api/ai-writer/templates/<int:template_id>', methods=['PUT'])
def api_update_template(template_id):
    if not verify_passcode():
        return jsonify({"error": "Chưa mở khóa hoặc Mã khóa không chính xác."}), 401
        
    name = request.form.get('name', '').strip()
    content = request.form.get('content', '').strip()
    
    # Hỗ trợ fallback nhận dữ liệu JSON
    if not name and not content:
        data = request.json or {}
        name = data.get('name', '').strip()
        content = data.get('content', '').strip()
        
    if not name:
        return jsonify({"error": "Tên mẫu hợp đồng không được bỏ trống."}), 400
        
    # Xử lý file docx đính kèm nếu có
    if 'file' in request.files:
        file = request.files['file']
        if file and file.filename != '':
            if not file.filename.lower().endswith('.docx'):
                return jsonify({"error": "Vui lòng đính kèm file Word (.docx)."}), 400
            try:
                temp_dir = tempfile.gettempdir()
                temp_path = os.path.join(temp_dir, file.filename)
                file.save(temp_path)
                
                doc = docx.Document(temp_path)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                content = '\n'.join(full_text)
                
                if os.path.exists(temp_path):
                    os.remove(temp_path)
            except Exception as e:
                return jsonify({"error": f"Lỗi đọc file docx: {str(e)}"}), 500
                
    if not content:
        return jsonify({"error": "Nội dung mẫu hợp đồng không được bỏ trống (hoặc phải đính kèm file docx)."}), 400
        
    success, err = database.update_template(template_id, name, content)
    if err:
        return jsonify({"error": err}), 400
    if not success:
        return jsonify({"error": "Không tìm thấy mẫu hợp đồng để cập nhật."}), 404
        
    return jsonify({"success": True, "message": "Cập nhật mẫu hợp đồng thành công."})

@app.route('/api/ai-writer/templates/<int:template_id>', methods=['DELETE'])
def api_delete_template(template_id):
    if not verify_passcode():
        return jsonify({"error": "Chưa mở khóa hoặc Mã khóa không chính xác."}), 401
    success = database.delete_template(template_id)
    if not success:
        return jsonify({"error": "Không tìm thấy mẫu hợp đồng để xóa."}), 404
    return jsonify({"success": True, "message": "Xóa mẫu hợp đồng thành công."})

@app.route('/api/ai-writer/chat', methods=['POST'])
def api_ai_writer_chat():
    if not verify_passcode():
        return jsonify({"error": "Chưa mở khóa hoặc Mã khóa không chính xác."}), 401
        
    if not GEMINI_API_KEY:
        return jsonify({"error": "Chưa cấu hình Claude API Key. Vui lòng cập nhật trong phần cài đặt."}), 400
        
    data = request.json or {}
    messages = data.get('messages', [])
    template_id = data.get('template_id')
    
    # Đọc template từ database nếu có
    template_context = ""
    if template_id:
        try:
            tmpl = database.get_template_by_id(int(template_id))
            if tmpl:
                template_context = tmpl.get('content', '')
        except Exception as e:
            print(f"[AI Writer] Error fetching template from DB: {e}")
                
    # Xây dựng Prompt hệ thống và ngữ cảnh
    system_instruction = """Bạn là trợ lý ảo chuyên soạn thảo hợp đồng pháp lý Việt Nam.
Nhiệm vụ của bạn là trò chuyện với người dùng để giúp họ soạn thảo và chỉnh sửa bản thảo hợp đồng.
Hãy phân tích lịch sử trò chuyện và yêu cầu mới nhất của người dùng.

YÊU CẦU ĐỊNH DẠNG PHẢN HỒI:
Trả về phản hồi dưới dạng JSON duy nhất chứa hai trường sau (không viết gì thêm ngoài JSON):
{
  "chat_response": "Lời thoại trò chuyện thân thiện, ngắn gọn thông báo tiến trình hoặc đặt câu hỏi bổ sung cho người dùng.",
  "contract_draft": "Toàn bộ văn bản hợp đồng được cập nhật hoặc tạo mới hoàn chỉnh dựa trên yêu cầu của người dùng, phân chia tiêu đề, điều khoản rõ ràng."
}

LƯU Ý SOẠN THẢO:
1. Trình bày văn bản hợp đồng chuẩn mực quốc gia Việt Nam (tiêu ngữ, tên hợp đồng viết hoa, căn lề, các ĐIỀU KHOẢN rõ ràng).
2. Nếu có mẫu hợp đồng đính kèm ở dưới, hãy tuân thủ cấu trúc của mẫu hợp đồng đó và điền hoặc bổ sung các thông tin còn thiếu dựa trên chat.
3. Nếu người dùng yêu cầu sửa đổi điều khoản, hãy cập nhật lại toàn bộ văn bản hợp đồng mới vào trường "contract_draft" (không chỉ viết đoạn sửa đổi).
"""
    if template_context:
        system_instruction += f"\n\nMẪU HỢP ĐỒNG THAM CHIẾU:\n{template_context}\n"
        
    # Tạo payload tin nhắn cho API AI
    api_messages = [{"role": "system", "content": system_instruction}]
    for msg in messages:
        role = msg.get('role', 'user')
        content = msg.get('content', '')
        if role in ['user', 'assistant']:
            api_messages.append({"role": role, "content": content})
            
    try:
        resp = requests.post(
            "https://api.vilao.ai/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {GEMINI_API_KEY}",
                "Content-Type": "application/json",
            },
            json={
                "model": "claude-opus-4-8",
                "messages": api_messages,
                "stream": False,
            },
            timeout=120,
        )
        resp.raise_for_status()
        resp_data = resp.json()
        resp_text = resp_data["choices"][0]["message"]["content"]
        
        # Parse JSON từ phản hồi của AI
        if "```json" in resp_text:
            resp_text = resp_text.split("```json")[1].split("```")[0]
        elif "```" in resp_text:
            resp_text = resp_text.split("```")[1].split("```")[0]
        else:
            start = resp_text.find('{')
            end = resp_text.rfind('}')
            if start != -1 and end != -1:
                resp_text = resp_text[start:end+1]
                
        parsed = json.loads(resp_text.strip())
        return jsonify({"success": True, "data": parsed})
    except Exception as e:
        return jsonify({"error": f"Lỗi gọi AI hoặc phân tích JSON: {str(e)}"}), 500

@app.route('/api/ai-writer/download', methods=['POST'])
def api_download_docx():
    import io
    import docx
    data = request.json or {}
    contract_text = data.get('contract_text', '')
    if not contract_text:
        return jsonify({"error": "Nội dung hợp đồng trống."}), 400
        
    doc = docx.Document()
    
    # Thiết lập lề trang
    for section in doc.sections:
        section.top_margin = docx.shared.Inches(0.79)
        section.bottom_margin = docx.shared.Inches(0.79)
        section.left_margin = docx.shared.Inches(1.18)
        section.right_margin = docx.shared.Inches(0.79)
        
    # Thiết lập font chữ Normal
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = docx.shared.Pt(13)
    
    import re
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def prevent_table_borders(table):
        tblPr = table._tbl.tblPr
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            tblBorders.append(border)
        tblPr.append(tblBorders)

    lines = contract_text.split('\n')
    in_header = True
    header_end_prefixes = ('căn cứ', 'hôm nay', 'bên a', 'bên b', 'điều', 'đại diện', 'mã số thuế', 'địa chỉ')
    current_table = None
    
    for line in lines:
        trimmed = line.strip()
        
        # Kiểm tra xem dòng này có phải là phần ký tên 2 cột không
        sig_match = re.match(r'^(.+?)(?:\s{4,}|\t+)(.+)$', line)
        if sig_match:
            in_header = False
            part1 = sig_match.group(1).strip()
            part2 = sig_match.group(2).strip()
            
            if not current_table:
                current_table = doc.add_table(rows=0, cols=2)
                prevent_table_borders(current_table)
                
            row = current_table.add_row()
            row.cells[0].width = docx.shared.Inches(3.2)
            row.cells[1].width = docx.shared.Inches(3.2)
            
            # Cột bên trái
            p_a = row.cells[0].paragraphs[0]
            p_a.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            run_a = p_a.add_run(part1)
            is_bold1 = part1.isupper() or part1.startswith('ĐẠI DIỆN') or part1.startswith('Đại diện')
            if is_bold1:
                run_a.bold = True
                
            # Cột bên phải
            p_b = row.cells[1].paragraphs[0]
            p_b.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            run_b = p_b.add_run(part2)
            is_bold2 = part2.isupper() or part2.startswith('ĐẠI DIỆN') or part2.startswith('Đại diện')
            if is_bold2:
                run_b.bold = True
                
        elif trimmed == "" and current_table is not None:
            # Giữ bảng mở và thêm dòng trống để tạo khoảng cách ký tên
            row = current_table.add_row()
            row.cells[0].width = docx.shared.Inches(3.2)
            row.cells[1].width = docx.shared.Inches(3.2)
            row.cells[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            row.cells[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            
        else:
            # Đóng bảng ký tên nếu đang mở
            current_table = None
            
            is_separator = len(trimmed) > 0 and len(trimmed) <= 30 and all(c in '-_* ' for c in trimmed)
            
            if in_header and not is_separator:
                trimmed_lower = trimmed.lower()
                if any(trimmed_lower.startswith(pref) for pref in header_end_prefixes) or trimmed.startswith('-') or trimmed.startswith('•'):
                    in_header = False
                    
            p = doc.add_paragraph()
            is_center = in_header and len(trimmed) > 0
            is_bold = trimmed.startswith('CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM') or trimmed.startswith('Độc lập - Tự do - Hạnh phúc') or trimmed.startswith('HỢP ĐỒNG') or trimmed.startswith('ĐIỀU') or trimmed.startswith('BÊN')
            
            if is_center:
                p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(trimmed)
                if is_bold:
                    run.bold = True
            else:
                if is_bold:
                    run = p.add_run(trimmed)
                    run.bold = True
                else:
                    p.add_run(line)
            
    # Lưu vào buffer
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return send_file(
        file_stream,
        as_attachment=True,
        download_name="Hop_Dong_Soan_Thao_AI.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# Khởi chạy luồng Telegram (hỗ trợ cả khi chạy qua Gunicorn/Waitress)
threading.Thread(target=telegram_worker, daemon=True).start()


if __name__ == '__main__':
    # Đọc port từ biến môi trường (Electron truyền vào), mặc định 5000
    port = int(os.environ.get('FLASK_PORT', 5000))
    host = os.environ.get('FLASK_HOST', '127.0.0.1')
    
    # Khởi chạy ngrok ở một luồng riêng để không chặn Flask
    threading.Thread(target=start_ngrok, args=(port,), daemon=True).start()
    
    # Khởi chạy Flask Server (tắt debug mode để chạy background threads ổn định hơn và tránh chạy 2 lần start_ngrok do reloader)
    app.run(host=host, port=port, debug=False)

